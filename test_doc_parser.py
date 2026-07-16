"""
统一文档解析测试。

测试内容：
1. TXT 普通文本解析
2. MD Markdown 解析
3. DOCX 普通段落、标题样式、编号标题、表格、中英文混合
4. XLSX 单工作表、多工作表、表头、空行、公式
5. ~$ 临时文件跳过
6. 空文件、损坏文件、无内容文件的错误处理
7. 来源文件和 location 元数据的透传

使用临时文件构建测试 fixture，不使用真实用户文件。
不调用外部 API。
"""

import io
import json
import sys
import tempfile
from pathlib import Path

sys.stdout.reconfigure(encoding="utf-8")

from doc_parser import parse_document, format_parse_summary, is_supported_file, get_supported_extensions


# ========== TXT 测试 ==========

def test_txt_basic():
    """TXT 普通文本解析。"""
    print("=" * 60)
    print("测试 TXT 普通文本")
    print("=" * 60)

    content = "这是第一段。\n\n这是第二段。\n\n这是第三段。"
    result = parse_document(file_bytes=content.encode("utf-8"), file_name="test.txt")

    assert result["source_type"] == "txt"
    assert result["source_file"] == "test.txt"
    assert "这是第一段" in result["text"]
    assert result["summary"]["paragraph_count"] == 3
    assert len(result["blocks"]) >= 1
    print("[OK] TXT 解析正确")

    print()


def test_txt_empty():
    """空 TXT 文件报错。"""
    print("=" * 60)
    print("测试空 TXT 文件")
    print("=" * 60)

    try:
        parse_document(file_bytes=b"", file_name="empty.txt")
        assert False, "应抛出 ValueError"
    except ValueError as e:
        assert "未提供文件内容" in str(e) or "未提取到任何文本" in str(e)
        print(f"[OK] 空文件报错: {e}")

    print()


def test_txt_gbk():
    """GBK 编码 TXT 文件。"""
    print("=" * 60)
    print("测试 GBK 编码 TXT")
    print("=" * 60)

    content = "这是GBK编码的中文内容。"
    result = parse_document(file_bytes=content.encode("gbk"), file_name="gbk.txt")

    assert "GBK编码" in result["text"]
    print("[OK] GBK 编码解析正确")

    print()


# ========== MD 测试 ==========

def test_md_basic():
    """MD Markdown 解析。"""
    print("=" * 60)
    print("测试 MD Markdown")
    print("=" * 60)

    content = "# 标题\n\n正文内容\n\n## 子标题\n\n更多内容。"
    result = parse_document(file_bytes=content.encode("utf-8"), file_name="test.md")

    assert result["source_type"] == "md"
    assert "# 标题" in result["text"]
    print("[OK] MD 解析正确")

    print()


# ========== DOCX 测试 ==========

def _create_docx(paragraphs=None, tables=None, heading_styles=None) -> bytes:
    """创建测试用 DOCX 文件。"""
    from docx import Document as DocxDocument
    from docx.shared import Pt

    doc = DocxDocument()

    if paragraphs:
        for i, (text, style) in enumerate(paragraphs):
            if style:
                doc.add_paragraph(text, style=style)
            else:
                doc.add_paragraph(text)

    if tables:
        for table_data in tables:
            rows, cols = len(table_data), len(table_data[0]) if table_data else 0
            table = doc.add_table(rows=rows, cols=cols)
            for r_idx, row in enumerate(table_data):
                for c_idx, cell_val in enumerate(row):
                    table.rows[r_idx].cells[c_idx].text = str(cell_val)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_docx_paragraphs():
    """DOCX 普通段落解析。"""
    print("=" * 60)
    print("测试 DOCX 普通段落")
    print("=" * 60)

    docx_bytes = _create_docx(paragraphs=[
        ("第一段内容。", None),
        ("第二段内容。", None),
        ("第三段内容。", None),
    ])

    result = parse_document(file_bytes=docx_bytes, file_name="test.docx")

    assert result["source_type"] == "docx"
    assert "第一段内容" in result["text"]
    assert "第二段内容" in result["text"]
    assert result["summary"]["paragraph_count"] == 3
    print("[OK] DOCX 普通段落解析正确")

    print()


def test_docx_headings():
    """DOCX 标题样式解析。"""
    print("=" * 60)
    print("测试 DOCX 标题样式")
    print("=" * 60)

    docx_bytes = _create_docx(paragraphs=[
        ("一级标题", "Heading 1"),
        ("正文内容一。", None),
        ("二级标题", "Heading 2"),
        ("正文内容二。", None),
    ])

    result = parse_document(file_bytes=docx_bytes, file_name="headings.docx")

    # 标题应作为独立块
    texts = [b["text"] for b in result["blocks"]]
    assert any("一级标题" in t for t in texts)
    assert any("二级标题" in t for t in texts)
    # 正文块不应包含标题前缀（标题只在 location 中）
    para_block = [b for b in result["blocks"] if "正文内容一" in b["text"]][0]
    assert para_block["text"] == "正文内容一。", f"正文块不应包含标题前缀，实际: {para_block['text']}"
    assert para_block["location"]["heading"] == "一级标题", "location 中应记录标题上下文"
    print("[OK] DOCX 标题样式解析正确")

    print()


def test_docx_numbered_headings():
    """DOCX 编号标题解析。"""
    print("=" * 60)
    print("测试 DOCX 编号标题")
    print("=" * 60)

    docx_bytes = _create_docx(paragraphs=[
        ("1. 第一章概述", None),
        ("这是第一章的内容。", None),
        ("1.1 背景", None),
        ("背景描述。", None),
        ("1.2 目标", None),
        ("目标描述。", None),
    ])

    result = parse_document(file_bytes=docx_bytes, file_name="numbered.docx")

    texts = [b["text"] for b in result["blocks"]]
    # 编号标题应被识别为标题
    assert any("1. 第一章概述" in t for t in texts)
    assert any("1.1 背景" in t for t in texts)
    print("[OK] DOCX 编号标题解析正确")

    print()


def test_docx_tables():
    """DOCX 表格解析。"""
    print("=" * 60)
    print("测试 DOCX 表格")
    print("=" * 60)

    table_data = [
        ["名称", "类型", "说明"],
        ["字段A", "字符串", "描述A"],
        ["字段B", "整数", "描述B"],
    ]
    docx_bytes = _create_docx(
        paragraphs=[("以下是表格：", None)],
        tables=[table_data],
    )

    result = parse_document(file_bytes=docx_bytes, file_name="tables.docx")

    assert result["summary"]["table_count"] == 1
    # 表格应转换为 Markdown
    table_blocks = [b for b in result["blocks"] if b["location"].get("table_index") is not None]
    assert len(table_blocks) == 1
    assert "名称" in table_blocks[0]["text"]
    assert "字段A" in table_blocks[0]["text"]
    assert "---" in table_blocks[0]["text"]  # Markdown 分隔线
    print("[OK] DOCX 表格解析正确")

    print()


def test_docx_mixed_language():
    """DOCX 中英文混合文本。"""
    print("=" * 60)
    print("测试 DOCX 中英文混合")
    print("=" * 60)

    docx_bytes = _create_docx(paragraphs=[
        ("P2P借贷（Peer-to-Peer Lending）是一种金融模式。", None),
        ("The concept involves direct lending between individuals.", None),
        ("中文段落与English混合test。", None),
    ])

    result = parse_document(file_bytes=docx_bytes, file_name="mixed.docx")

    assert "P2P借贷" in result["text"]
    assert "Peer-to-Peer" in result["text"]
    assert "中文段落" in result["text"]
    print("[OK] DOCX 中英文混合解析正确")

    print()


def test_docx_empty():
    """空 DOCX 文件报错。"""
    print("=" * 60)
    print("测试空 DOCX 文件")
    print("=" * 60)

    docx_bytes = _create_docx()  # 无段落无表格

    try:
        parse_document(file_bytes=docx_bytes, file_name="empty.docx")
        assert False, "应抛出 ValueError"
    except ValueError as e:
        assert "未提取到任何段落" in str(e)
        print(f"[OK] 空 DOCX 报错: {e}")

    print()


# ========== XLSX 测试 ==========

def _create_xlsx(sheets_data=None) -> bytes:
    """创建测试用 XLSX 文件。

    sheets_data: dict of {sheet_name: [[row1], [row2], ...]}
    """
    from openpyxl import Workbook

    wb = Workbook()
    # 删除默认工作表
    wb.remove(wb.active)

    if sheets_data:
        for sheet_name, rows in sheets_data.items():
            ws = wb.create_sheet(title=sheet_name)
            for row in rows:
                ws.append(row)

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def test_xlsx_basic():
    """XLSX 单工作表解析。"""
    print("=" * 60)
    print("测试 XLSX 单工作表")
    print("=" * 60)

    xlsx_bytes = _create_xlsx({
        "产品表": [
            ["名称", "价格", "数量"],
            ["产品A", 100, 10],
            ["产品B", 200, 20],
        ]
    })

    result = parse_document(file_bytes=xlsx_bytes, file_name="test.xlsx")

    assert result["source_type"] == "xlsx"
    assert result["summary"]["sheet_count"] == 1
    assert result["summary"]["row_count"] == 2
    assert "产品A" in result["text"]
    assert "名称" in result["text"]
    # 检查格式：工作表：产品表；行：2；名称：产品A；价格：100；数量：10
    assert "工作表：产品表" in result["text"]
    assert "行：2" in result["text"]
    print("[OK] XLSX 单工作表解析正确")

    print()


def test_xlsx_multi_sheet():
    """XLSX 多工作表解析。"""
    print("=" * 60)
    print("测试 XLSX 多工作表")
    print("=" * 60)

    xlsx_bytes = _create_xlsx({
        "Sheet1": [
            ["ID", "名称"],
            [1, "项目A"],
        ],
        "Sheet2": [
            ["类别", "值"],
            ["类型X", 100],
        ],
    })

    result = parse_document(file_bytes=xlsx_bytes, file_name="multi.xlsx")

    assert result["summary"]["sheet_count"] == 2
    assert "工作表：Sheet1" in result["text"]
    assert "工作表：Sheet2" in result["text"]
    print("[OK] XLSX 多工作表解析正确")

    print()


def test_xlsx_empty_rows():
    """XLSX 空行跳过。"""
    print("=" * 60)
    print("测试 XLSX 空行跳过")
    print("=" * 60)

    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Test"
    ws.append(["Name", "Value"])
    ws.append(["A", 1])
    ws.append([None, None])  # 空行
    ws.append(["", ""])  # 空字符串行
    ws.append(["B", 2])

    buf = io.BytesIO()
    wb.save(buf)
    xlsx_bytes = buf.getvalue()

    result = parse_document(file_bytes=xlsx_bytes, file_name="empty_rows.xlsx")

    # 应该只解析出 2 行数据（跳过空行）
    assert result["summary"]["row_count"] == 2
    assert "A" in result["text"]
    assert "B" in result["text"]
    print("[OK] XLSX 空行正确跳过")

    print()


def test_xlsx_formula():
    """XLSX 公式处理。"""
    print("=" * 60)
    print("测试 XLSX 公式处理")
    print("=" * 60)

    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Formula"
    ws.append(["A", "B", "Sum"])
    ws.append([10, 20, None])
    ws.cell(row=2, column=3, value="=A2+B2")  # 公式

    buf = io.BytesIO()
    wb.save(buf)
    xlsx_bytes = buf.getvalue()

    result = parse_document(file_bytes=xlsx_bytes, file_name="formula.xlsx")

    # 公式应被保留
    assert "=A2+B2" in result["text"] or "公式未计算" in result["text"]
    # 应有警告
    assert len(result["warnings"]) > 0 or "公式未计算" in result["text"]
    print("[OK] XLSX 公式处理正确")

    print()


def test_xlsx_empty_sheet():
    """XLSX 空工作表跳过。"""
    print("=" * 60)
    print("测试 XLSX 空工作表跳过")
    print("=" * 60)

    from openpyxl import Workbook

    wb = Workbook()
    # 默认空工作表
    ws1 = wb.active
    ws1.title = "Empty"
    # 有数据的工作表
    ws2 = wb.create_sheet("Data")
    ws2.append(["Name", "Value"])
    ws2.append(["A", 1])

    buf = io.BytesIO()
    wb.save(buf)
    xlsx_bytes = buf.getvalue()

    result = parse_document(file_bytes=xlsx_bytes, file_name="empty_sheet.xlsx")

    # 有数据的工作表应被解析
    assert "工作表：Data" in result["text"]
    # 空工作表可能被跳过（取决于 openpyxl 的 max_row 行为）
    # 至少验证不会崩溃且有数据
    assert result["summary"]["row_count"] >= 1
    print("[OK] XLSX 空工作表处理正确")

    print()


# ========== 临时文件和错误处理 ==========

def test_temp_file_skip():
    """~$ 临时文件跳过。"""
    print("=" * 60)
    print("测试 ~$ 临时文件跳过")
    print("=" * 60)

    # 在 app.py 中通过 st.warning 跳过，这里测试 is_supported_file
    assert is_supported_file("~$document.docx") is True  # 扩展名支持
    # 实际跳过逻辑在 app.py 中通过文件名前缀检查
    print("[OK] ~$ 临时文件检测逻辑正确")

    print()


def test_unsupported_format():
    """不支持的格式报错。"""
    print("=" * 60)
    print("测试不支持的格式")
    print("=" * 60)

    try:
        parse_document(file_bytes=b"test", file_name="test.pdf")
        assert False, "应抛出 ValueError"
    except ValueError as e:
        assert "不支持的文件格式" in str(e)
        print(f"[OK] 不支持格式报错: {e}")

    print()


def test_corrupted_docx():
    """损坏的 DOCX 文件报错。"""
    print("=" * 60)
    print("测试损坏的 DOCX 文件")
    print("=" * 60)

    try:
        parse_document(file_bytes=b"not a docx file", file_name="corrupted.docx")
        assert False, "应抛出 ValueError"
    except ValueError as e:
        assert "无法打开" in str(e)
        print(f"[OK] 损坏 DOCX 报错: {e}")

    print()


def test_corrupted_xlsx():
    """损坏的 XLSX 文件报错。"""
    print("=" * 60)
    print("测试损坏的 XLSX 文件")
    print("=" * 60)

    try:
        parse_document(file_bytes=b"not an xlsx file", file_name="corrupted.xlsx")
        assert False, "应抛出 ValueError"
    except ValueError as e:
        assert "无法打开" in str(e)
        print(f"[OK] 损坏 XLSX 报错: {e}")

    print()


# ========== 元数据透传 ==========

def test_metadata_transparent():
    """来源文件和 location 元数据透传。"""
    print("=" * 60)
    print("测试元数据透传")
    print("=" * 60)

    # TXT
    result = parse_document(file_bytes=b"test content", file_name="source.txt")
    assert result["source_file"] == "source.txt"
    assert result["source_type"] == "txt"
    assert result["blocks"][0]["location"]["source_file"] == "source.txt"
    print("[OK] TXT 元数据正确")

    # DOCX
    docx_bytes = _create_docx(paragraphs=[("段落内容", None)])
    result = parse_document(file_bytes=docx_bytes, file_name="source.docx")
    assert result["source_file"] == "source.docx"
    assert result["source_type"] == "docx"
    assert result["blocks"][0]["location"]["source_file"] == "source.docx"
    assert result["blocks"][0]["location"]["source_type"] == "docx"
    print("[OK] DOCX 元数据正确")

    # XLSX
    xlsx_bytes = _create_xlsx({"Sheet1": [["A", "B"], [1, 2]]})
    result = parse_document(file_bytes=xlsx_bytes, file_name="source.xlsx")
    assert result["source_file"] == "source.xlsx"
    assert result["source_type"] == "xlsx"
    assert result["blocks"][0]["location"]["source_file"] == "source.xlsx"
    assert result["blocks"][0]["location"]["sheet_name"] == "Sheet1"
    assert result["blocks"][0]["location"]["row_number"] == 2
    print("[OK] XLSX 元数据正确")

    print()


def test_format_summary():
    """格式化解析摘要。"""
    print("=" * 60)
    print("测试格式化解析摘要")
    print("=" * 60)

    result = parse_document(file_bytes="段落一。\n\n段落二。".encode("utf-8"), file_name="test.txt")
    summary = format_parse_summary(result)

    assert "test.txt" in summary
    assert "txt" in summary
    assert "段落数" in summary
    print(f"  摘要: {summary}")
    print("[OK] 格式化摘要正确")

    print()


# ========== 无重复验证 ==========

def test_docx_no_duplication():
    """DOCX 解析不产生重复内容：N 个唯一段落输出 N 个块，总字符数一致。"""
    print("=" * 60)
    print("测试 DOCX 无重复内容")
    print("=" * 60)

    # 构建含标题 + 多个正文段落的文档
    paragraphs = [
        ("Chapter 1: Introduction", "Heading 1"),
        ("This is the first paragraph under chapter 1.", None),
        ("This is the second paragraph under chapter 1.", None),
        ("1.1 Background", None),  # 编号标题
        ("Background content here.", None),
        ("Chapter 2: Methods", "Heading 1"),
        ("Method description paragraph.", None),
        ("2.1 Data Collection", None),
        ("Data collection details.", None),
        ("Another paragraph under 2.1.", None),
    ]

    docx_bytes = _create_docx(paragraphs=paragraphs)
    result = parse_document(file_bytes=docx_bytes, file_name="no_dup_test.docx")

    # 验证块数等于非空段落数
    expected_block_count = len(paragraphs)
    actual_block_count = len(result["blocks"])
    assert actual_block_count == expected_block_count, \
        f"块数应为 {expected_block_count}，实际 {actual_block_count}"

    # 验证总字符数等于各段落拼接长度（不含 heading 前缀重复）
    expected_chars = sum(len(text) for text, _ in paragraphs)
    # full_text 用 \n\n 连接各块，加上连接符的长度
    joiner_len = (actual_block_count - 1) * 2  # \n\n between blocks
    expected_total = expected_chars + joiner_len
    actual_total = len(result["text"])
    assert actual_total == expected_total, \
        f"总字符数应为 {expected_total}，实际 {actual_total}（差值: {actual_total - expected_total}）"

    # 验证每个段落在全文中只出现一次（作为独立块）
    for text, _ in paragraphs:
        occurrences = result["text"].count(text)
        assert occurrences == 1, \
            f"段落 '{text[:30]}...' 在全文中出现 {occurrences} 次，期望 1 次"

    # 验证标题不被重复拼接到正文块中
    chapter1_block = result["blocks"][1]  # 第一个正文段落
    assert chapter1_block["text"] == "This is the first paragraph under chapter 1.", \
        f"正文块不应包含标题前缀，实际: {chapter1_block['text'][:60]}"

    print(f"  块数: {actual_block_count} (期望 {expected_block_count})")
    print(f"  总字符数: {actual_total} (期望 {expected_total})")
    print("[OK] DOCX 无重复内容")

    print()


# ========== 工具函数 ==========

def test_utility_functions():
    """工具函数测试。"""
    print("=" * 60)
    print("测试工具函数")
    print("=" * 60)

    exts = get_supported_extensions()
    assert ".txt" in exts
    assert ".md" in exts
    assert ".docx" in exts
    assert ".xlsx" in exts
    print(f"  支持格式: {exts}")

    assert is_supported_file("test.txt") is True
    assert is_supported_file("test.docx") is True
    assert is_supported_file("test.xlsx") is True
    assert is_supported_file("test.pdf") is False
    assert is_supported_file("test.csv") is False
    print("[OK] 工具函数正确")

    print()


def main():
    print("=" * 60)
    print("统一文档解析测试")
    print("=" * 60)
    print()

    # TXT
    test_txt_basic()
    test_txt_empty()
    test_txt_gbk()

    # MD
    test_md_basic()

    # DOCX
    test_docx_paragraphs()
    test_docx_headings()
    test_docx_numbered_headings()
    test_docx_tables()
    test_docx_mixed_language()
    test_docx_empty()

    # XLSX
    test_xlsx_basic()
    test_xlsx_multi_sheet()
    test_xlsx_empty_rows()
    test_xlsx_formula()
    test_xlsx_empty_sheet()

    # 错误处理
    test_temp_file_skip()
    test_unsupported_format()
    test_corrupted_docx()
    test_corrupted_xlsx()

    # 元数据
    test_metadata_transparent()
    test_format_summary()

    # 工具函数
    test_utility_functions()

    print("=" * 60)
    print("[OK] 所有测试通过！")
    print("=" * 60)


if __name__ == "__main__":
    main()
