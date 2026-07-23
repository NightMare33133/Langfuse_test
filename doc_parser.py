"""
统一源文档解析模块。

支持格式：.txt, .md, .docx, .xlsx
输出：规范化文本 + 来源元数据（source_file, source_type, location）

设计原则：
- 后续切分和题目生成逻辑只消费规范化文本，不关心原始格式
- 每个文本块附带 location 元数据，方便追溯来源
- 不修改原始文件
"""

import re
from pathlib import Path


# ========== 统一解析接口 ==========

def parse_document(file_path: str = None, file_bytes: bytes = None, file_name: str = "") -> dict:
    """统一文档解析入口。

    Args:
        file_path: 文件路径（与 file_bytes 二选一）
        file_bytes: 文件字节内容（Streamlit UploadedFile.getvalue()）
        file_name: 文件名（用于推断格式和元数据）

    Returns:
        dict: {
            "text": str,              # 合并后的规范化全文
            "blocks": list[dict],     # 文本块列表，每块含 text + location
            "source_file": str,       # 来源文件名
            "source_type": str,       # txt / md / docx / xlsx
            "summary": dict,          # 解析摘要（段落数、表格数等）
            "warnings": list[str],    # 解析警告
        }

    Raises:
        ValueError: 不支持的格式、空文件、无可提取内容
    """
    # 确定文件名和字节内容
    if file_path:
        p = Path(file_path)
        file_name = file_name or p.name
        file_bytes = p.read_bytes()
    if not file_bytes:
        raise ValueError("未提供文件内容")

    # 推断格式
    ext = Path(file_name).suffix.lower()
    if ext not in (".txt", ".md", ".docx", ".xlsx", ".xls", ".csv"):
        raise ValueError(f"不支持的文件格式: {ext}（支持 .txt, .md, .docx, .xlsx, .xls, .csv）")

    source_type = ext.lstrip(".")

    # 分派解析
    if source_type in ("txt", "md"):
        result = _parse_text(file_bytes, file_name, source_type)
    elif source_type == "docx":
        result = _parse_docx(file_bytes, file_name)
    elif source_type == "xlsx":
        result = _parse_xlsx(file_bytes, file_name)
    elif source_type == "csv":
        result = _parse_csv(file_bytes, file_name)
    elif source_type == "xls":
        result = _parse_xls(file_bytes, file_name)
    else:
        raise ValueError(f"未实现的解析器: {source_type}")

    result["source_file"] = file_name
    result["source_type"] = source_type

    # 验证非空
    if not result["text"].strip():
        raise ValueError(f"文件 {file_name} 中未提取到任何文本内容")

    return result


# ========== TXT / MD 解析 ==========

def _parse_text(file_bytes: bytes, file_name: str, source_type: str) -> dict:
    """解析纯文本或 Markdown 文件。"""
    # 尝试 UTF-8，回退 GBK
    try:
        text = file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        text = file_bytes.decode("gbk", errors="replace")

    blocks = [{"text": text, "location": {"source_file": file_name, "source_type": source_type}}]
    paragraph_count = len([b for b in text.split("\n\n") if b.strip()])

    return {
        "text": text,
        "blocks": blocks,
        "summary": {"paragraph_count": paragraph_count, "table_count": 0},
        "warnings": [],
    }


# ========== DOCX 解析 ==========

def _parse_docx(file_bytes: bytes, file_name: str) -> dict:
    """解析 Word .docx 文件。"""
    import io
    from docx import Document as DocxDocument

    warnings = []

    try:
        doc = DocxDocument(io.BytesIO(file_bytes))
    except Exception as e:
        raise ValueError(f"无法打开 Word 文件: {e}")

    blocks = []
    paragraph_count = 0
    table_count = 0
    current_heading = ""

    # 解析段落
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # 检测标题
        style_name = (para.style.name or "").lower()
        is_heading = "heading" in style_name

        # 检测编号标题（如 1.、1.1、1.2.3）
        numbered_heading = _detect_numbered_heading(text)

        if is_heading or numbered_heading:
            current_heading = text
            blocks.append({
                "text": text,
                "location": {
                    "source_file": file_name,
                    "source_type": "docx",
                    "heading": text,
                    "style": para.style.name if is_heading else "numbered",
                },
            })
        else:
            # 普通段落：只存原文，标题上下文记录在 location 中
            blocks.append({
                "text": text,
                "location": {
                    "source_file": file_name,
                    "source_type": "docx",
                    "heading": current_heading or None,
                    "paragraph_index": paragraph_count,
                },
            })
        paragraph_count += 1

    # 解析表格
    for table_idx, table in enumerate(doc.tables):
        table_text = _table_to_markdown(table, file_name, table_idx)
        if table_text.strip():
            blocks.append({
                "text": table_text,
                "location": {
                    "source_file": file_name,
                    "source_type": "docx",
                    "table_index": table_idx,
                    "heading": current_heading or None,
                },
            })
            table_count += 1

    if not blocks:
        raise ValueError(f"Word 文件 {file_name} 中未提取到任何段落或表格")

    # 合并全文
    full_text = "\n\n".join(b["text"] for b in blocks)

    return {
        "text": full_text,
        "blocks": blocks,
        "summary": {"paragraph_count": paragraph_count, "table_count": table_count},
        "warnings": warnings,
    }


def _detect_numbered_heading(text: str) -> bool:
    """检测编号标题，如 1.、1.1、1.2.3、第X章 等。"""
    patterns = [
        r"^\d+\.\s+\S",        # 1. 标题
        r"^\d+\.\d+\s+\S",    # 1.1 标题
        r"^\d+\.\d+\.\d+",    # 1.2.3 标题
        r"^第[一二三四五六七八九十\d]+[章节篇]",  # 第X章
    ]
    for p in patterns:
        if re.match(p, text):
            return True
    return False


def _table_to_markdown(table, file_name: str, table_idx: int) -> str:
    """将 Word 表格转换为 Markdown 格式。"""
    rows = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            cell_text = cell.text.strip().replace("\n", " ")
            cells.append(cell_text)
        if any(c for c in cells):  # 跳过纯空行
            rows.append(cells)

    if not rows:
        return ""

    # 构建 Markdown 表格
    lines = []
    # 表头
    header = rows[0]
    lines.append("| " + " | ".join(header) + " |")
    lines.append("| " + " | ".join(["---"] * len(header)) + " |")
    # 数据行
    for row in rows[1:]:
        # 补齐列数
        while len(row) < len(header):
            row.append("")
        lines.append("| " + " | ".join(row[:len(header)]) + " |")

    return "\n".join(lines)


# ========== XLSX 解析 ==========

def _parse_xlsx(file_bytes: bytes, file_name: str) -> dict:
    """解析 Excel .xlsx 文件。"""
    import io
    from openpyxl import load_workbook

    warnings = []

    try:
        wb = load_workbook(io.BytesIO(file_bytes), data_only=False)
    except Exception as e:
        raise ValueError(f"无法打开 Excel 文件: {e}")

    blocks = []
    sheet_count = 0
    row_count = 0

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        if ws.max_row is None or ws.max_row < 1:
            warnings.append(f"工作表「{sheet_name}」为空，已跳过")
            continue

        # 读取表头
        headers = []
        for col in range(1, (ws.max_column or 0) + 1):
            val = ws.cell(row=1, column=col).value
            headers.append(str(val).strip() if val is not None else f"列{col}")

        # 读取数据行
        sheet_row_count = 0
        for row_idx in range(2, (ws.max_row or 0) + 1):
            # 检查是否为空行
            row_values = []
            is_empty = True
            for col in range(1, (ws.max_column or 0) + 1):
                cell = ws.cell(row=row_idx, column=col)
                val = cell.value
                if val is not None and str(val).strip():
                    is_empty = False
                row_values.append(val)

            if is_empty:
                continue

            # 构建可读文本
            parts = [f"工作表：{sheet_name}；行：{row_idx}"]
            for col_idx, (header, val) in enumerate(zip(headers, row_values)):
                if val is None:
                    continue
                cell = ws.cell(row=row_idx, column=col_idx + 1)
                display_val = _format_cell_value(cell, warnings, sheet_name, row_idx, header)
                parts.append(f"{header}：{display_val}")

            block_text = "；".join(parts)
            blocks.append({
                "text": block_text,
                "location": {
                    "source_file": file_name,
                    "source_type": "xlsx",
                    "sheet_name": sheet_name,
                    "row_number": row_idx,
                    "headers": headers,
                },
            })
            sheet_row_count += 1

        if sheet_row_count > 0:
            sheet_count += 1
            row_count += sheet_row_count

    if not blocks:
        raise ValueError(f"Excel 文件 {file_name} 中未提取到任何数据")

    # 合并全文
    full_text = "\n\n".join(b["text"] for b in blocks)

    return {
        "text": full_text,
        "blocks": blocks,
        "summary": {"sheet_count": sheet_count, "row_count": row_count, "table_count": sheet_count},
        "warnings": warnings,
    }


def _format_cell_value(cell, warnings: list, sheet_name: str, row_idx: int, header: str) -> str:
    """格式化单元格值，处理公式和缓存值。"""
    val = cell.value
    if val is None:
        return ""

    val_str = str(val)

    # 检查是否为公式
    if val_str.startswith("="):
        # 尝试获取缓存计算值
        # openpyxl data_only=False 时保留公式，需要重新打开 data_only=True 获取缓存值
        # 这里我们标记公式并保留公式本身
        cached = None
        try:
            # data_only 模式下可以读取缓存值，但我们当前是 data_only=False
            # 所以只能标记公式
            pass
        except Exception:
            pass

        if cached is not None:
            return f"{val_str} (缓存值: {cached})"
        else:
            warnings.append(f"工作表「{sheet_name}」行{row_idx}「{header}」含公式未计算: {val_str[:50]}")
            return f"{val_str} [公式未计算]"

    return val_str


# ========== CSV 解析 ==========

def _parse_csv(file_bytes: bytes, file_name: str) -> dict:
    """解析 CSV 文件。支持 UTF-8、UTF-8 BOM、GBK 编码探测。"""
    import io
    import pandas as pd

    warnings = []

    # 编码探测
    encoding = None
    for enc in ("utf-8-sig", "utf-8", "gbk"):
        try:
            file_bytes.decode(enc)
            encoding = enc
            break
        except (UnicodeDecodeError, LookupError):
            continue
    if encoding is None:
        try:
            import charset_normalizer
            result = charset_normalizer.from_bytes(file_bytes).best()
            encoding = result.encoding if result else "utf-8"
        except Exception:
            encoding = "utf-8"
        warnings.append(f"自动检测编码: {encoding}")

    try:
        df = pd.read_csv(io.BytesIO(file_bytes), encoding=encoding, dtype=str, keep_default_na=False)
    except Exception as e:
        raise ValueError(f"无法解析 CSV 文件: {e}")

    if df.empty:
        raise ValueError(f"CSV 文件 {file_name} 为空或无有效数据")

    headers = [str(c) for c in df.columns.tolist()]
    blocks = []
    row_count = 0

    for row_idx, row in enumerate(df.values.tolist(), start=2):
        row_values = [str(v) for v in row]
        if not any(v.strip() for v in row_values):
            continue

        parts = [f"行：{row_idx}"]
        for header, val in zip(headers, row_values):
            if val.strip():
                parts.append(f"{header}：{val}")

        block_text = "；".join(parts)
        blocks.append({
            "text": block_text,
            "location": {
                "source_file": file_name,
                "source_type": "csv",
                "sheet_name": "CSV",
                "row_number": row_idx,
                "headers": headers,
            },
        })
        row_count += 1

    if not blocks:
        raise ValueError(f"CSV 文件 {file_name} 中未提取到任何数据")

    full_text = "\n\n".join(b["text"] for b in blocks)

    return {
        "text": full_text,
        "blocks": blocks,
        "summary": {"sheet_count": 1, "row_count": row_count, "table_count": 1},
        "warnings": warnings,
    }


# ========== XLS 解析 ==========

def _parse_xls(file_bytes: bytes, file_name: str) -> dict:
    """解析 Excel .xls 文件。使用 pandas + xlrd。"""
    import io
    import pandas as pd

    warnings = []

    try:
        sheets = pd.read_excel(io.BytesIO(file_bytes), sheet_name=None, engine='xlrd', dtype=str)
    except ImportError:
        raise ValueError("XLS 格式需要安装 xlrd 库。请运行: pip install xlrd")
    except Exception as e:
        raise ValueError(f"无法打开 XLS 文件: {e}")

    warnings.append("XLS 格式不保留公式信息，单元格显示值可能为缓存计算结果")

    blocks = []
    sheet_count = 0
    row_count = 0

    for sheet_name, df in sheets.items():
        if df.empty:
            warnings.append(f"工作表「{sheet_name}」为空，已跳过")
            continue

        df = df.fillna("")
        headers = [str(c) for c in df.columns.tolist()]

        sheet_row_count = 0
        for row_idx, row in enumerate(df.values.tolist(), start=2):
            row_values = [str(v) for v in row]
            if not any(v.strip() for v in row_values):
                continue

            parts = [f"工作表：{sheet_name}；行：{row_idx}"]
            for header, val in zip(headers, row_values):
                if val.strip():
                    parts.append(f"{header}：{val}")

            block_text = "；".join(parts)
            blocks.append({
                "text": block_text,
                "location": {
                    "source_file": file_name,
                    "source_type": "xls",
                    "sheet_name": str(sheet_name),
                    "row_number": row_idx,
                    "headers": headers,
                },
            })
            sheet_row_count += 1

        if sheet_row_count > 0:
            sheet_count += 1
            row_count += sheet_row_count

    if not blocks:
        raise ValueError(f"XLS 文件 {file_name} 中未提取到任何数据")

    full_text = "\n\n".join(b["text"] for b in blocks)

    return {
        "text": full_text,
        "blocks": blocks,
        "summary": {"sheet_count": sheet_count, "row_count": row_count, "table_count": sheet_count},
        "warnings": warnings,
    }


# ========== 工具函数 ==========

def get_supported_extensions() -> list:
    """返回支持的文件扩展名列表。"""
    return [".txt", ".md", ".docx", ".xlsx", ".xls", ".csv"]


def is_supported_file(file_name: str) -> bool:
    """检查文件是否为支持的格式。"""
    return Path(file_name).suffix.lower() in get_supported_extensions()


def format_parse_summary(result: dict) -> str:
    """格式化解析摘要为可读文本。"""
    parts = []
    parts.append(f"文件: {result['source_file']}")
    parts.append(f"格式: {result['source_type']}")

    summary = result.get("summary", {})
    if "paragraph_count" in summary:
        parts.append(f"段落数: {summary['paragraph_count']}")
    if "table_count" in summary and summary["table_count"]:
        parts.append(f"表格数: {summary['table_count']}")
    if "sheet_count" in summary and summary["sheet_count"]:
        parts.append(f"工作表数: {summary['sheet_count']}")
    if "row_count" in summary and summary["row_count"]:
        parts.append(f"数据行数: {summary['row_count']}")

    blocks = result.get("blocks", [])
    parts.append(f"文本块数: {len(blocks)}")
    parts.append(f"总字符数: {len(result.get('text', ''))}")

    warnings = result.get("warnings", [])
    if warnings:
        parts.append(f"警告: {len(warnings)} 条")

    return " | ".join(parts)
